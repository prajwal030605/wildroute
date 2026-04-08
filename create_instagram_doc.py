from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Helpers ──
def add_heading(text, level=1, color=(29, 158, 117)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.bold = True
        run.font.name = "Arial"
    return p

def add_para(text="", bold=False, italic=False, color=None, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1D9E75')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_label(text):
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(29, 158, 117)
    # Shade background
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0F2A1E')
    rPr.append(shd)
    return p

GREEN = (29, 158, 117)
WHITE = (255, 255, 255)
DARK  = (30, 30, 30)
GRAY  = (100, 100, 100)

# ════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════
add_para()
add_para()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WildRoute")
run.font.size = Pt(32)
run.font.bold = True
run.font.name = "Arial"
run.font.color.rgb = RGBColor(*GREEN)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Instagram Content — 7-Day Brand Awareness Pack")
r2.font.size = Pt(14)
r2.font.name = "Arial"
r2.font.color.rgb = RGBColor(*GRAY)

add_para()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("@go.wildroute  ·  Educational & Brand Awareness  ·  Week 1")
r3.font.size = Pt(11)
r3.font.name = "Arial"
r3.font.color.rgb = RGBColor(*GRAY)
r3.italic = True

add_para()
add_divider()
add_para()

# ════════════════════════════════════════
# OVERVIEW TABLE
# ════════════════════════════════════════
add_heading("Content Overview", level=1)
add_para()

table = doc.add_table(rows=8, cols=4)
table.style = "Table Grid"
headers = ["Day", "Post Type", "Theme", "Goal"]
hrow = table.rows[0]
for i, h in enumerate(headers):
    cell = hrow.cells[i]
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1D9E75')
    cell.paragraphs[0]._p.get_or_add_pPr().append(shd)

rows_data = [
    ["Day 1", "Single Image", "Brand Introduction", "Awareness"],
    ["Day 2", "Carousel (5 slides)", "How to Pick an Adventure Agency", "Education"],
    ["Day 3", "Single Image", "Destination Spotlight — Hampta Pass", "Inspiration"],
    ["Day 4", "Carousel (4 slides)", "What's Actually Included in a Trek?", "Education"],
    ["Day 5", "Single Image / Reel", "The Problem WildRoute Solves", "Awareness"],
    ["Day 6", "Poll Story + Post", "Which adventure is on your list?", "Engagement"],
    ["Day 7", "Single Image", "Join the Waitlist — Launch Coming Soon", "Conversion"],
]
for i, row_data in enumerate(rows_data):
    row = table.rows[i+1]
    fill = "0F2A1E" if i % 2 == 0 else "111111"
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(val)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(200, 200, 200) if j > 0 else RGBColor(29, 158, 117)
        if j == 0:
            run.font.bold = True
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        cell.paragraphs[0]._p.get_or_add_pPr().append(shd)

add_para()
add_divider()

# ════════════════════════════════════════
posts = [
    {
        "day": "Day 1",
        "type": "Single Image",
        "theme": "Brand Introduction",
        "visual": "Clean dark background with WildRoute logo centered. Tagline: 'Your route to real adventure.' Subtle mountain silhouette at the bottom. Brand green (#1D9E75) accents.",
        "caption": """🏔️ Something big is coming for India's adventure community.

We've all been there — you want to trek Hampta Pass or go rafting in Rishikesh, but you have no idea which agency to trust. Prices are all over the place. Reviews feel fake. You book and hope for the best.

That ends with WildRoute.

WildRoute is India's first platform to compare and book verified adventure agencies — all in one place. Trekking, rafting, paragliding, bungee jumping, and more.

✅ Verified agencies (GST, permits & insurance checked)
✅ Real reviews from real trekkers
✅ Transparent pricing — no hidden costs
✅ Secure booking with full refund protection

We're launching soon. Drop your email at the link in bio to get early access 🎒

—
#WildRoute #AdventureIndia #TrekkingIndia #IndianAdventure #HimalayanTreks #AdventureTravel #ComingSoon #Trekking #Rafting #Paragliding #IndiaTravel #WildRouteIndia""",
        "hashtags": "#WildRoute #AdventureIndia #TrekkingIndia #IndianAdventure #HimalayanTreks #AdventureTravel #ComingSoon #Trekking #Rafting #Paragliding #IndiaTravel #WildRouteIndia",
        "cta": "Drop your email at the link in bio to get early access 🎒"
    },
    {
        "day": "Day 2",
        "type": "Carousel — 5 Slides",
        "theme": "How to Pick a Legit Adventure Agency",
        "visual": "Slide 1: Bold hook text on dark background. Slides 2-5: Each tip on its own card with icon + short explanation. Slide 5: CTA slide with WildRoute branding.",
        "caption": """Most trekking scams are 100% avoidable 🧵

Before you pay any adventure agency, check these 5 things 👇
(Swipe to see each one)

Slide 1 — 🚩 Are they GST registered?
Any legit business running commercial tours must be GST registered. Ask for their GSTIN before paying. If they can't provide it, walk away.

Slide 2 — 🛡️ Do they have proper permits?
Every trek in restricted zones (Ladakh, parts of Uttarakhand, Northeast) needs permits. Ask your agency if permits are included and what they cover.

Slide 3 — 🏥 Is insurance included?
Altitude sickness, accidents, evacuation — adventure goes wrong sometimes. Your agency should include trek insurance for every group member.

Slide 4 — ⭐ Are the reviews verified?
Google reviews can be faked. Ask the agency for past client contacts, or look for platforms where reviews are only from verified bookings.

Slide 5 — 📋 Is the itinerary detailed?
Vague itineraries = hidden costs later. A good agency gives you day-by-day plans with meal details, stay types, and what's NOT included.

WildRoute verifies all of this for every agency on our platform 🟢
Link in bio to join the waitlist.

—
#TrekkingTips #AdventureTips #HowToTrek #TrekkingIndia #AdventureIndia #SafeAdventure #WildRoute #TravelTips #IndianTrekker #MountainLife""",
        "hashtags": "#TrekkingTips #AdventureTips #HowToTrek #TrekkingIndia #AdventureIndia #SafeAdventure #WildRoute #TravelTips #IndianTrekker #MountainLife",
        "cta": "Save this before your next trek 📌"
    },
    {
        "day": "Day 3",
        "type": "Single Image",
        "theme": "Destination Spotlight — Hampta Pass",
        "visual": "Stunning wide shot of Hampta Pass (mountains, snow, valley contrast). Dark overlay with location name + key stats overlaid. WildRoute logo watermark bottom right.",
        "caption": """📍 Hampta Pass, Himachal Pradesh — 14,100 ft

This is what happens when two worlds collide.

On one side: the lush green Kullu Valley.
On the other: the barren, moonscape of Spiti.

Hampta Pass sits at 14,100 ft and connects two completely different worlds in a single 5-day trek. It's one of the most dramatic transitions you'll ever walk through in the Himalayas.

Best for: First-time high altitude trekkers
Season: June – September
Difficulty: Moderate
Starting point: Manali

🏔️ What makes it special:
→ The contrast between lush and barren landscapes is unlike anything else in India
→ You end at Chandrataal Lake — one of the most beautiful lakes in Asia
→ Fully doable without prior trekking experience

Bookmark this for your 2025 adventure list 🔖

Which trek is on your radar? Tell us in the comments 👇

—
#HamtaPass #TrekkingIndia #HimachalPradesh #Manali #HimalayanTreks #MountainTrekking #IndianHimalayas #AdventureTravel #WildRoute #TrekLife #Chandrataal #Spiti""",
        "hashtags": "#HamtaPass #TrekkingIndia #HimachalPradesh #Manali #HimalayanTreks #MountainTrekking #IndianHimalayas #AdventureTravel #WildRoute #TrekLife #Chandrataal #Spiti",
        "cta": "Bookmark this for your 2025 adventure list 🔖"
    },
    {
        "day": "Day 4",
        "type": "Carousel — 4 Slides",
        "theme": "What's Actually Included in a Trek Package?",
        "visual": "Slide 1: Question hook — 'You're paying ₹8,000 for a trek. But what does that actually include?' Slides 2-4: Checklist format. Final slide: WildRoute shows you exactly what's included before you pay.",
        "caption": """You're paying ₹8,000 for a trek. Do you actually know what you're getting? 👇

Swipe through the full checklist ➡️

Most people book a trek without knowing what's included — and end up surprised at base camp when they're told food costs extra.

Here's what a GOOD trek package should always include:

🏕️ Accommodation
Tents on the trail + guesthouse/hotel before/after. Ask specifically: shared or private? Sleeping bag included?

🍽️ Meals
Most legit packages include all meals from Day 1 dinner to last day lunch. Watch out for "meals not included" in the fine print.

🎒 Equipment
Trekking poles, crampons (if needed), and safety gear should be provided or at minimum available to rent from the agency.

👨‍🏫 Guide & Porter
A certified mountain guide is non-negotiable. A porter is often optional — but clarify the ratio (how many people per guide?).

🚌 Transport
Most packages include transport from the nearest city (Manali, Leh, Rishikesh). Confirm pickup/drop point.

📋 Permits
If the trek requires permits (Ladakh, Northeast), they should be included in the price — not charged separately.

At WildRoute, every listing shows exactly what's included. No guessing. No surprises.
Link in bio 👆

—
#TrekkingIndia #TrekTips #AdventureTips #WildRoute #HimalayanTreks #TravelIndia #MountainLife #IndianAdventure #TrekkerLife #AdventureReady""",
        "hashtags": "#TrekkingIndia #TrekTips #AdventureTips #WildRoute #HimalayanTreks #TravelIndia #MountainLife #IndianAdventure #TrekkerLife #AdventureReady",
        "cta": "Save this — read it before you book your next trek 📌"
    },
    {
        "day": "Day 5",
        "type": "Single Image / Short Reel",
        "theme": "The Problem WildRoute Solves",
        "visual": "Split graphic: Left side shows chaos (multiple tabs, WhatsApp messages, confusing prices). Right side shows WildRoute's clean comparison interface. Bold headline: 'There had to be a better way.'",
        "caption": """Booking an adventure in India shouldn't feel like this 😮‍💨

You want to go trekking. Simple enough, right?

But then it starts:
→ 3 hours googling agencies
→ 10 WhatsApp messages to get a price
→ 4 different quotes for the same trek
→ Zero idea which agency is actually legit
→ Pray and pay, hoping it works out

This is how 99% of people book adventure in India right now.

There's no single place to compare agencies, check real reviews, see transparent pricing, or book safely.

That's the exact problem WildRoute was built to solve.

One platform. All verified agencies. Real reviews. Transparent prices. Safe booking.

We're building it right now — and we want you on the list when we launch 🚀

Link in bio → Join the waitlist (it's free)

—
#WildRoute #AdventureIndia #TrekkingIndia #StartupIndia #AdventureBooking #TravelProblem #IndianTravel #TrekLife #AdventureTravel #BuildingInPublic #IndianStartup #WildRouteIndia""",
        "hashtags": "#WildRoute #AdventureIndia #TrekkingIndia #StartupIndia #AdventureBooking #TravelProblem #IndianTravel #TrekLife #AdventureTravel #BuildingInPublic #IndianStartup #WildRouteIndia",
        "cta": "Link in bio → Join the waitlist (it's free)"
    },
    {
        "day": "Day 6",
        "type": "Engagement Post + Story Poll",
        "theme": "Community Engagement — Which Adventure?",
        "visual": "Grid of 4 adventure photos (trekking, rafting, paragliding, bungee). Clean layout with activity names. Caption-driven engagement. Story: Poll between options.",
        "caption": """If you could do ONE of these this weekend, what would it be? 👇

🥾 A. Trekking — pack your bag and disappear into the mountains
🚣 B. Rafting — battle the rapids with your crew
🪂 C. Paragliding — take the leap and fly over a valley
🪢 D. Bungee jumping — fall and feel everything

Drop your answer in the comments — and tell us WHERE in India you'd do it 📍

(We'll be sharing the best suggestions in our stories 👀)

—
For us, there's nothing like waking up at 5am with your tent frosted over, knowing you've got a high pass to cross before noon. But we might be biased 🏔️

—
Follow @go.wildroute for destination guides, trekking tips, and everything adventure 🎒

#AdventureIndia #TrekkingIndia #Rafting #Paragliding #BungeeJumping #WhatsYourAdventure #WildRoute #MountainLife #IndianAdventure #TravelIndia #WeekendAdventure #AdventureLovers""",
        "hashtags": "#AdventureIndia #TrekkingIndia #Rafting #Paragliding #BungeeJumping #WhatsYourAdventure #WildRoute #MountainLife #IndianAdventure #TravelIndia #WeekendAdventure #AdventureLovers",
        "cta": "Drop your answer in the comments 👇",
        "story_note": "Story Poll: 'Which adventure are you doing next?' Options: 🥾 Trekking vs 🪂 Paragliding | Follow-up story: 'Which feels more YOU?' — 🚣 Rafting vs 🪢 Bungee"
    },
    {
        "day": "Day 7",
        "type": "Single Image — CTA",
        "theme": "Join the Waitlist",
        "visual": "Bold, clean design. WildRoute logo + 'Launching Soon' badge. Text: 'India's first platform to compare & book adventure agencies.' Green CTA button graphic: 'Join the waitlist →'. Trekker silhouette on ridge.",
        "caption": """India's adventure scene is about to change 🏔️

WildRoute is launching soon — and we're giving early access to everyone who signs up before we go live.

Here's what you get as an early trekker:
→ First access to verified agencies before they fill up
→ Exclusive first-mover discounts at launch
→ A say in what we build next

And if you're an adventure agency:
→ List for FREE
→ Zero commission for your first 6 months
→ Direct access to thousands of trekkers from day one

We've already got {trekkerCount}+ trekkers and {agencyCount}+ agencies on the list.

Don't be the last one in 🎒

👉 Link in bio — takes 10 seconds, costs nothing.

—
#WildRoute #AdventureIndia #TrekkingIndia #JoinTheWaitlist #LaunchingSoon #IndianAdventure #HimalayanTreks #AdventureAgency #TrekkingCommunity #WildRouteIndia #EarlyAccess #AdventureBooking""",
        "hashtags": "#WildRoute #AdventureIndia #TrekkingIndia #JoinTheWaitlist #LaunchingSoon #IndianAdventure #HimalayanTreks #AdventureAgency #TrekkingCommunity #WildRouteIndia #EarlyAccess #AdventureBooking",
        "cta": "Link in bio → Join the waitlist (10 seconds, free)"
    },
]

# ════════════════════════════════════════
# POSTS
# ════════════════════════════════════════
for i, post in enumerate(posts):
    doc.add_page_break()

    # Day header
    p = doc.add_paragraph()
    run = p.add_run(f"  {post['day']}  ")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1D9E75')
    rPr.append(shd)

    # Theme title
    h = doc.add_heading(post["theme"], level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(29, 158, 117)
        run.font.name = "Arial"

    # Post type
    p = doc.add_paragraph()
    run = p.add_run(f"Format: {post['type']}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*GRAY)

    add_para()

    # Visual direction
    doc.add_heading("🎨 Visual Direction", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(*GRAY)
        run.font.name = "Arial"
    p = doc.add_paragraph(post["visual"])
    p.paragraph_format.left_indent = Inches(0.2)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(150, 150, 150)
        run.font.italic = True

    add_para()

    # Caption
    doc.add_heading("✍️ Caption", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(*GRAY)
        run.font.name = "Arial"

    # Split caption into paragraphs
    lines = post["caption"].strip().split("\n")
    for line in lines:
        p = doc.add_paragraph(line if line else "")
        p.paragraph_format.left_indent = Inches(0.2)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(30, 30, 30)
        if not line:
            p.paragraph_format.space_after = Pt(4)

    add_para()

    # CTA
    doc.add_heading("📣 CTA", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(*GRAY)
        run.font.name = "Arial"
    p = doc.add_paragraph(post["cta"])
    p.paragraph_format.left_indent = Inches(0.2)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(29, 158, 117)

    # Hashtags
    add_para()
    doc.add_heading("# Hashtags", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(*GRAY)
        run.font.name = "Arial"
    p = doc.add_paragraph(post["hashtags"])
    p.paragraph_format.left_indent = Inches(0.2)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 120, 200)

    # Story note if exists
    if "story_note" in post:
        add_para()
        doc.add_heading("📱 Story Idea", level=2)
        for run in doc.paragraphs[-1].runs:
            run.font.color.rgb = RGBColor(*GRAY)
            run.font.name = "Arial"
        p = doc.add_paragraph(post["story_note"])
        p.paragraph_format.left_indent = Inches(0.2)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(150, 100, 200)
            run.font.italic = True

    add_divider()

# Save
out = "/sessions/upbeat-laughing-curie/mnt/wildroute/WildRoute_Instagram_Content_Week1.docx"
doc.save(out)
print(f"Saved: {out}")
