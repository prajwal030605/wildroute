from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

GREEN  = RGBColor(29, 158, 117)
WHITE  = RGBColor(255, 255, 255)
GRAY   = RGBColor(120, 120, 120)
LGRAY  = RGBColor(180, 180, 180)
DARK   = RGBColor(30, 30, 30)
YELLOW = RGBColor(245, 158, 11)
RED    = RGBColor(239, 68, 68)

def shade_para(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def divider(color="1D9E75"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1');    b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)

def h1(text, color=GREEN):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = color; r.font.name = "Arial"
    return p

def h2(text, color=GRAY):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = color; r.font.name = "Arial"
    return p

def body(text, color=DARK, size=11, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(size)
    r.font.color.rgb = color; r.bold = bold; r.italic = italic
    return p

def tag(text, fill="1D9E75", txt_color=WHITE):
    p = doc.add_paragraph()
    r = p.add_run(f"  {text}  ")
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.bold = True
    r.font.color.rgb = txt_color
    rPr = r._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill); rPr.append(shd)
    return p

def space(): doc.add_paragraph()

def slide_box(num, headline, subline, body_lines, icon, tip_label, do_text, dont_text, why_text):
    """Render one carousel slide as a nicely formatted block."""
    # Slide number tag
    tag(f"SLIDE {num} of 6", fill="111111", txt_color=RGBColor(100,100,100))

    # Icon + headline
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f"{icon}  ")
    r1.font.size = Pt(22); r1.font.name = "Arial"
    r2 = p.add_run(headline)
    r2.font.size = Pt(18); r2.font.bold = True
    r2.font.name = "Arial"; r2.font.color.rgb = WHITE

    # Subline
    p2 = doc.add_paragraph(subline)
    p2.paragraph_format.left_indent = Inches(0.2)
    for r in p2.runs:
        r.font.name = "Arial"; r.font.size = Pt(11)
        r.font.color.rgb = LGRAY; r.italic = True

    space()

    # Tip label
    tag(f"  {tip_label}  ", fill="0F2A1E", txt_color=GREEN)

    # Body lines
    for line in body_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.3)
        for r in p.runs:
            r.font.name = "Arial"; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(60,60,60)

    space()

    # Do / Don't table
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    col_w = [4500, 4500]

    headers = [("✅  DO THIS", "2D3A2E"), ("❌  AVOID THIS", "3A2D2D")]
    vals    = [(do_text, "1D9E75"), (dont_text, "EF4444")]

    for col_i, (hdr, fill) in enumerate(headers):
        cell = tbl.rows[0].cells[col_i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(hdr)
        r.font.name = "Arial"; r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = WHITE
        pPr = cell.paragraphs[0]._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
        pPr.append(shd)

    for col_i, (val, col_hex) in enumerate(vals):
        cell = tbl.rows[1].cells[col_i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(val)
        r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(*bytes.fromhex(col_hex))
        cell.paragraphs[0]._p.get_or_add_pPr()

    space()

    # Why it matters
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run("Why this matters:  ")
    r1.font.name = "Arial"; r1.font.size = Pt(10)
    r1.font.bold = True; r1.font.color.rgb = YELLOW
    r2 = p.add_run(why_text)
    r2.font.name = "Arial"; r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY; r2.italic = True


# ════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("WildRoute Instagram")
r.font.size = Pt(11); r.font.name = "Arial"; r.font.color.rgb = GRAY; r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Day 2 — Carousel Post")
r.font.size = Pt(28); r.font.bold = True; r.font.name = "Arial"
r.font.color.rgb = GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("How to Pick a Legit Adventure Agency")
r.font.size = Pt(16); r.font.name = "Arial"; r.font.color.rgb = LGRAY

space()
divider()
space()

# ════════════════════════════════════════════════════════════
# POST OVERVIEW
# ════════════════════════════════════════════════════════════
h1("Post Overview")

tbl = doc.add_table(rows=6, cols=2)
tbl.style = "Table Grid"
rows = [
    ("Format",     "Instagram Carousel — 6 Slides"),
    ("Theme",      "Educational — How to pick a legit adventure agency"),
    ("Goal",       "Build trust, educate audience, position WildRoute as the solution"),
    ("Tone",       "Helpful, direct, slightly urgent — like advice from a fellow trekker"),
    ("Best time",  "Tuesday or Wednesday, 7–9 PM IST"),
    ("CTA",        "Save this post + Follow @go.wildroute"),
]
for i, (k, v) in enumerate(rows):
    fill = "0F0F0F" if i % 2 == 0 else "141414"
    for col_i, text in enumerate([k, v]):
        cell = tbl.rows[i].cells[col_i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.bold = (col_i == 0)
        r.font.color.rgb = GREEN if col_i == 0 else LGRAY
        pPr = cell.paragraphs[0]._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
        pPr.append(shd)

space()
divider()
space()

# ════════════════════════════════════════════════════════════
# SLIDES
# ════════════════════════════════════════════════════════════
h1("Slide-by-Slide Breakdown")
space()

# ── SLIDE 1 — HOOK ──
tag("SLIDE 1 of 6  ·  HOOK SLIDE", fill="1D9E75")
space()

h2("Headline (bold, large)")
body("Most trekking scams are 100% avoidable.", color=WHITE, size=20, bold=True)
space()

h2("Subheadline")
body("5 things to check before you pay any adventure agency.", color=LGRAY, size=13, italic=True)
space()

h2("Visual Direction")
body("Dark background (#0a0a0a). WildRoute logo top-right corner, small.", color=GRAY, size=10, italic=True, indent=0.2)
body("Large bold white text centered. Subtle red warning icon top-left.", color=GRAY, size=10, italic=True, indent=0.2)
body("Bottom text: 'Swipe to see each one →' in green.", color=GRAY, size=10, italic=True, indent=0.2)
space()

h2("Design Notes")
body("Font: Bold, clean sans-serif. Text takes up 70% of slide.", color=GRAY, size=10, italic=True, indent=0.2)
body("Use 🚩 emoji prominently — it signals 'red flag tips' and stops the scroll.", color=GRAY, size=10, italic=True, indent=0.2)

space()
divider(color="333333")
space()

# ── SLIDES 2–6 ──
slides = [
    {
        "num": 2, "icon": "🏛️",
        "headline": "Are they GST registered?",
        "subline": "The most basic trust signal — and most people never ask.",
        "tip_label": "TIP 1 of 5",
        "body_lines": [
            "Any business running commercial adventure tours in India must be GST registered.",
            "A GSTIN proves they're a legitimate, tax-paying entity — not just a WhatsApp number.",
            "Ask for it before you transfer any money. Legit agencies will share it instantly.",
        ],
        "do": "Ask for their GSTIN upfront and verify it at gst.gov.in",
        "dont": "Pay anyone who can't immediately provide their GST number",
        "why": "Unregistered operators have no legal accountability. If something goes wrong, you have no recourse.",
    },
    {
        "num": 3, "icon": "📋",
        "headline": "Do they have the right permits?",
        "subline": "Trekking without permits = illegal. In some zones, it's also dangerous.",
        "tip_label": "TIP 2 of 5",
        "body_lines": [
            "Restricted zones — Ladakh, Northeast, parts of Uttarakhand — require Inner Line Permits.",
            "Forest treks need separate forest department permits.",
            "Your agency must arrange these. If they say 'you'll sort it yourself', that's a red flag.",
        ],
        "do": "Confirm permits are included in the package price and ask to see them before departure",
        "dont": "Assume permits are included — always ask specifically what's covered",
        "why": "Getting caught without permits means being turned back mid-trek or fined. Don't risk it.",
    },
    {
        "num": 4, "icon": "🏥",
        "headline": "Is trek insurance included?",
        "subline": "Altitude sickness, injuries, evacuation — it happens more than you think.",
        "tip_label": "TIP 3 of 5",
        "body_lines": [
            "High-altitude rescue can cost ₹1–5 lakh. Helicopter evacuation from Ladakh? Even more.",
            "Good agencies include insurance for every group member in the package price.",
            "Ask specifically: What does it cover? Medical? Evacuation? Accidental death?",
        ],
        "do": "Get the insurance policy number and check what's covered before you leave",
        "dont": "Accept vague assurances like 'don't worry, we'll handle it' without written proof",
        "why": "One accident without insurance can cost more than 10 treks. It's non-negotiable.",
    },
    {
        "num": 5, "icon": "⭐",
        "headline": "Are the reviews actually real?",
        "subline": "Google reviews can be bought for ₹5 each. Most people don't know this.",
        "tip_label": "TIP 4 of 5",
        "body_lines": [
            "Signs of fake reviews: generic praise, no specific details, sudden review bursts.",
            "Signs of real reviews: mentions guide names, specific route details, honest pros/cons.",
            "Best signal: Ask the agency for 2–3 past client contacts and actually call them.",
        ],
        "do": "Look for detailed, specific reviews — and ask for past client references directly",
        "dont": "Trust star ratings alone — a 4.9 with 200 reviews can still be entirely fake",
        "why": "Real reviews protect you from agencies that look great online but cut corners on the ground.",
    },
    {
        "num": 6, "icon": "📄",
        "headline": "Is the itinerary detailed enough?",
        "subline": "Vague itineraries are where hidden costs hide.",
        "tip_label": "TIP 5 of 5",
        "body_lines": [
            "A good itinerary gives you: day-by-day plans, altitude for each camp, meal details,",
            "accommodation type (tent/guesthouse), what's NOT included, and cancellation terms.",
            "If the itinerary is just '5 days trekking, meals and stay included' — walk away.",
        ],
        "do": "Ask for a full written itinerary + terms before paying even a deposit",
        "dont": "Book based on a verbal description or a vague PDF with no specific inclusions listed",
        "why": "Hidden costs discovered mid-trek ruin the experience and break trust. Clarity upfront saves everyone.",
    },
]

for s in slides:
    slide_box(
        num=s["num"], icon=s["icon"],
        headline=s["headline"], subline=s["subline"],
        body_lines=s["body_lines"], tip_label=s["tip_label"],
        do_text=s["do"], dont_text=s["dont"], why_text=s["why"],
    )
    space()
    divider(color="222222")
    space()

# ── SLIDE 6 — CTA ──
tag("SLIDE 6 of 6  ·  CTA SLIDE", fill="1D9E75")
space()
h2("Headline")
body("WildRoute verifies all of this — for every agency on our platform.", color=WHITE, size=16, bold=True)
space()
h2("Body text on slide")
body("GST checked. Permits verified. Insurance confirmed. Reviews validated.", color=LGRAY, size=12)
space()
h2("CTA on slide")
body("Join the waitlist → Link in bio", color=GREEN, size=13, bold=True)
space()
h2("Visual Direction")
body("Green background (#1D9E75). White WildRoute logo centered.", color=GRAY, size=10, italic=True, indent=0.2)
body("Checklist of 4 verified items with white tick icons.", color=GRAY, size=10, italic=True, indent=0.2)
body("'Launching soon' badge. Large CTA button graphic.", color=GRAY, size=10, italic=True, indent=0.2)

space()
divider()
space()

# ════════════════════════════════════════════════════════════
# FULL CAPTION
# ════════════════════════════════════════════════════════════
h1("Full Instagram Caption")
space()

caption_lines = [
    "Most trekking scams are 100% avoidable 🧵",
    "",
    "Before you pay any adventure agency, check these 5 things 👇",
    "(Swipe through each one ➡️)",
    "",
    "🚩 1. Are they GST registered?",
    "Any legit business running commercial tours must be GST registered. Ask for their GSTIN",
    "before paying. If they hesitate, walk away.",
    "",
    "📋 2. Do they have proper permits?",
    "Restricted zones like Ladakh and the Northeast require Inner Line Permits. Your agency",
    "should arrange these — not you. Confirm it's included before booking.",
    "",
    "🏥 3. Is insurance included?",
    "High-altitude rescue can cost ₹1–5 lakh. A good agency includes trek insurance for",
    "every member. Ask what it covers: medical, evacuation, accidental death.",
    "",
    "⭐ 4. Are the reviews actually real?",
    "Google reviews can be faked for ₹5 each. Look for specific details — guide names,",
    "route notes, honest cons. Or ask the agency for past client contacts and call them.",
    "",
    "📄 5. Is the itinerary detailed?",
    "Vague itineraries = hidden costs later. A proper itinerary shows day-by-day plans,",
    "what's included, what's NOT, and cancellation terms. If it's just '5 days trekking' — red flag.",
    "",
    "—",
    "WildRoute verifies all of this for every agency on our platform.",
    "Launching soon → Join the waitlist at the link in bio 🎒",
    "",
    "Save this before your next trek 📌",
    "",
    "—",
    "#TrekkingTips #AdventureTips #HowToTrek #TrekkingIndia #AdventureIndia",
    "#SafeAdventure #WildRoute #TravelTips #IndianTrekker #MountainLife",
    "#HimalayanTreks #TrekSmart #AdventureReady #TrekkerLife #IndianAdventure",
]

for line in caption_lines:
    p = doc.add_paragraph(line if line else "")
    p.paragraph_format.left_indent = Inches(0.2)
    if not line:
        p.paragraph_format.space_after = Pt(2)
        continue
    is_hashtag = line.startswith("#")
    is_tip     = any(line.startswith(x) for x in ["🚩","📋","🏥","⭐","📄"])
    is_dash    = line == "—"
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(10)
        if is_hashtag:
            r.font.color.rgb = RGBColor(80, 120, 200)
        elif is_tip:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(11)
        elif is_dash:
            r.font.color.rgb = GRAY
        else:
            r.font.color.rgb = RGBColor(50, 50, 50)

space()
divider()
space()

# ════════════════════════════════════════════════════════════
# POSTING CHECKLIST
# ════════════════════════════════════════════════════════════
h1("Pre-Post Checklist")
space()

checks = [
    ("Caption copy-pasted and proofread", "Make sure emojis render correctly on mobile"),
    ("All 6 slides designed and exported at 1080x1080px", "Square format works best for carousels"),
    ("First slide is scroll-stopping", "Test: would YOU stop scrolling for this?"),
    ("Hashtags added (not in caption body — in first comment)", "Keep caption clean, hashtags in comment"),
    ("Location tag added", "Tag a relevant location — e.g. 'Indian Himalayas'"),
    ("Story created to promote the carousel", "Share slide 1 to story with 'Swipe on feed' sticker"),
    ("Reply to ALL comments within first hour", "Early engagement boosts the algorithm significantly"),
]

tbl2 = doc.add_table(rows=len(checks)+1, cols=3)
tbl2.style = "Table Grid"

hdr_row = tbl2.rows[0]
for col_i, hdr in enumerate(["", "Task", "Note"]):
    cell = hdr_row.cells[col_i]
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(hdr)
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.bold = True
    r.font.color.rgb = WHITE
    pPr = cell.paragraphs[0]._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1D9E75')
    pPr.append(shd)

for i, (task, note) in enumerate(checks):
    row = tbl2.rows[i+1]
    fill = "0F0F0F" if i % 2 == 0 else "141414"
    for col_i, (text, col) in enumerate([("☐", GRAY), (task, LGRAY), (note, GRAY)]):
        cell = row.cells[col_i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.color.rgb = col
        if col_i == 1: r.font.bold = True
        pPr = cell.paragraphs[0]._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
        pPr.append(shd)

space()
divider()
space()

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("WildRoute  ·  @go.wildroute  ·  Day 2 Content Brief")
r.font.name = "Arial"; r.font.size = Pt(9); r.font.color.rgb = GRAY; r.italic = True

out = "/sessions/upbeat-laughing-curie/mnt/wildroute/WildRoute_Day2_Carousel.docx"
doc.save(out)
print(f"Saved: {out}")
